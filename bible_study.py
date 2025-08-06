import os
from crewai import Agent, Task, Crew, Process
from crewai_tools import SerperDevTool
from langchain_google_genai import ChatGoogleGenerativeAI

class BibleStudyAgents:
    """Initializes agents with the user-selected Gemini model and API key."""
    def __init__(self, model_name, api_key):
        self.llm = ChatGoogleGenerativeAI(
            model=model_name,
            google_api_key=api_key,
            temperature=0.5,
            convert_system_message_to_human=True
        )

    def biblical_historian(self):
        return Agent(
            role='Biblical Historian & Archaeologist',
            goal='Provide a comprehensive historical, cultural, and literary background for a given book of the Bible, in the specified language.',
            backstory="With a PhD from Jerusalem University and fluency in multiple languages, you provide the crucial context that makes the biblical text come alive.",
            llm=self.llm,
            allow_delegation=False,
            verbose=True
        )

    def exegetical_theologian(self):
        return Agent(
            role='Exegetical Theologian',
            goal='Analyze the text of a Bible book to uncover its main theological themes, key verses, and structure, presenting the findings in the specified language.',
            backstory="As a systematic theologian, you are an expert at exegesis—drawing out the intended meaning of the text for a global audience.",
            llm=self.llm,
            allow_delegation=False,
            verbose=True
        )
    
    def practical_application_guide(self):
        return Agent(
            role='Pastoral Guide & Counselor',
            goal='Create practical, thought-provoking application questions and prayer points based on the themes of a Bible book, written in the specified language.',
            backstory="You are a seasoned pastor skilled in multicultural ministry, crafting questions that bridge the gap between ancient text and modern life.",
            llm=self.llm,
            allow_delegation=False,
            verbose=True
        )

    def senior_editor(self):
        return Agent(
            role='Senior Editor for Christian Publishing',
            goal='Compile the work of the other agents into a single, cohesive, and beautifully formatted Bible study guide in the specified language.',
            backstory="You work for an international Christian publishing house, ensuring every manuscript is professional, theologically sound, and ready for a global readership.",
            llm=self.llm,
            allow_delegation=False,
            verbose=True
        )

class BibleStudyTasks:
    """Defines the tasks for creating the Bible study guide."""
    def historical_context_task(self, agent, bible_book, language):
        return Task(
            description=f"Create the 'Historical Background' section for a study guide on **{bible_book}**. Your output MUST be in {language}.",
            expected_output=f"A Markdown section on the historical background of {bible_book}, written entirely in {language}.",
            agent=agent
        )

    def theological_analysis_task(self, agent, bible_book, language):
        return Task(
            description=f"Create the 'Theological Themes & Key Verses' section for **{bible_book}**. Your output MUST be in {language}. Use a well-known {language} Bible translation for quotes.",
            expected_output=f"A detailed Markdown section on theological themes of {bible_book}, written entirely in {language}.",
            agent=agent
        )
    
    def application_task(self, agent, bible_book, language):
        return Task(
            description=f"Create the 'Practical Application & Reflection' section for **{bible_book}**. Your output MUST be in {language}.",
            expected_output=f"An encouraging Markdown section with discussion questions and prayer points for {bible_book}, written entirely in {language}.",
            agent=agent
        )

    def editing_task(self, agent, bible_book, language, context):
        return Task(
            description=f"Compile all sections into a single study guide. The final output must be in {language}. The main title should be the {language} translation for 'A Study Guide to the Book of {bible_book}'.",
            expected_output=f"A complete, well-formatted Markdown document in {language}.",
            agent=agent,
            context=context,
            output_file=f'final_study_guide_{language.lower()}.md'
        )

def create_bible_study_crew(bible_book, language, selected_model, gemini_api_key, serper_api_key):
    """
    This function initializes the AI crew with user-provided credentials and model selection.
    """
    agents = BibleStudyAgents(model_name=selected_model, api_key=gemini_api_key)
    tasks = BibleStudyTasks()

    # Correctly initialize the search tool with the user's key
    search_tool = SerperDevTool(api_key=serper_api_key)

    # Define Agents
    historian = agents.biblical_historian()
    theologian = agents.exegetical_theologian()
    pastor = agents.practical_application_guide()
    editor = agents.senior_editor()
    
    # Correctly assign the search tool to the agents that need it
    historian.tools = [search_tool]
    theologian.tools = [search_tool]
    
    # Define Tasks
    task1 = tasks.historical_context_task(historian, bible_book, language)
    task2 = tasks.theological_analysis_task(theologian, bible_book, language)
    task3 = tasks.application_task(pastor, bible_book, language)
    task4 = tasks.editing_task(editor, bible_book, language, [task1, task2, task3])
    
    return Crew(
        agents=[historian, theologian, pastor, editor],
        tasks=[task1, task2, task3, task4],
        process=Process.sequential,
        verbose=2
    )
    
 import streamlit as st
from bible_study_crew import create_bible_study_crew
import markdown_pdf
from docx import Document
import base64

# Helper functions for exporting remain the same
def markdown_to_pdf(md_content):
    pdf = markdown_pdf.MarkdownPdf(body_font_size=12)
    pdf.add_section(md_content, toc=False)
    return pdf.get_buffer()

def markdown_to_docx(md_content):
    doc = Document()
    for line in md_content.split('\n'):
        if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1)
        elif line.strip() == '---': doc.add_page_break()
        else: doc.add_paragraph(line)
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Bible book translations remain the same
ENGLISH_BOOKS = ["Genesis", "Exodus", "Leviticus", "Numbers", "Deuteronomy", "Joshua", "Judges", "Ruth", "1 Samuel", "2 Samuel", "1 Kings", "2 Kings", "1 Chronicles", "2 Chronicles", "Ezra", "Nehemiah", "Esther", "Job", "Psalms", "Proverbs", "Ecclesiastes", "Song of Solomon", "Isaiah", "Jeremiah", "Lamentations", "Ezekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadiah", "Jonah", "Micah", "Nahum", "Habakkuk", "Zephaniah", "Haggai", "Zechariah", "Malachi", "Matthew", "Mark", "Luke", "John", "Acts", "Romans", "1 Corinthians", "2 Corinthians", "Galatians", "Ephesians", "Philippians", "Colossians", "1 Thessalonians", "2 Thessalonians", "1 Timothy", "2 Timothy", "Titus", "Philemon", "Hebrews", "James", "1 Peter", "2 Peter", "1 John", "2 John", "3 John", "Jude", "Revelation"]
BIBLE_BOOKS_TRANSLATIONS = {
    "English": ENGLISH_BOOKS,
    "German": ["Genesis", "Exodus", "Levitikus", "Numeri", "Deuteronomium", "Josua", "Richter", "Ruth", "1. Samuel", "2. Samuel", "1. Könige", "2. Könige", "1. Chronik", "2. Chronik", "Esra", "Nehemia", "Esther", "Hiob", "Psalmen", "Sprüche", "Prediger", "Hohelied", "Jesaja", "Jeremia", "Klagelieder", "Hesekiel", "Daniel", "Hosea", "Joel", "Amos", "Obadja", "Jona", "Micha", "Nahum", "Habakuk", "Zefanja", "Haggai", "Sacharja", "Maleachi", "Matthäus", "Markus", "Lukas", "Johannes", "Apostelgeschichte", "Römer", "1. Korinther", "2. Korinther", "Galater", "Epheser", "Philipper", "Kolosser", "1. Thessalonicher", "2. Thessalonicher", "1. Timotheus", "2. Timotheus", "Titus", "Philemon", "Hebräer", "Jakobus", "1. Petrus", "2. Petrus", "1. Johannes", "2. Johannes", "3. Johannes", "Judas", "Offenbarung"],
    "French": ["Genèse", "Exode", "Lévitique", "Nombres", "Deutéronome", "Josué", "Juges", "Ruth", "1 Samuel", "2 Samuel", "1 Rois", "2 Rois", "1 Chroniques", "2 Chroniques", "Esdras", "Néhémie", "Esther", "Job", "Psaumes", "Proverbes", "Ecclésiaste", "Cantique des Cantiques", "Ésaïe", "Jérémie", "Lamentations", "Ézéchiel", "Daniel", "Osée", "Joël", "Amos", "Abdias", "Jonas", "Michée", "Nahum", "Habacuc", "Sophonie", "Aggée", "Zacharie", "Malachie", "Matthieu", "Marc", "Luc", "Jean", "Actes", "Romains", "1 Corinthiens", "2 Corinthiens", "Galates", "Éphésiens", "Philippiens", "Colossiens", "1 Thessaloniciens", "2 Thessaloniciens", "1 Timothée", "2 Timothée", "Tite", "Philémon", "Hébreux", "Jacques", "1 Pierre", "2 Pierre", "1 Jean", "2 Jean", "3 Jean", "Jude", "Apocalypse"],
    "Swahili": ["Mwanzo", "Kutoka", "Walawi", "Hesabu", "Kumbukumbu la Torati", "Yoshua", "Waamuzi", "Ruthu", "1 Samweli", "2 Samweli", "1 Wafalme", "2 Wafalme", "1 Mambo ya Nyakati", "2 Mambo ya Nyakati", "Ezra", "Nehemia", "Esta", "Ayubu", "Zaburi", "Methali", "Mhubiri", "Wimbo Ulio Bora", "Isaya", "Yeremia", "Maombolezo", "Ezekieli", "Danieli", "Hosea", "Yoeli", "Amosi", "Obadia", "Yona", "Mika", "Nahumu", "Habakuki", "Sefania", "Hagai", "Zekaria", "Malaki", "Mathayo", "Marko", "Luka", "Yohana", "Matendo", "Warumi", "1 Wakorintho", "2 Wakorintho", "Wagalatia", "Waefeso", "Wafilipi", "Wakolosai", "1 Wathesalonike", "2 Wathesalonike", "1 Timotheo", "2 Timotheo", "Tito", "Filemoni", "Waebrania", "Yakobo", "1 Petro", "2 Petro", "1 Yohana", "2 Yohana", "3 Yohana", "Yuda", "Ufunuo"]
}

# --- CORRECTED: List of specific Gemini Models ---
GEMINI_MODEL_LIST = [
    "gemini/gemini-2.5-pro-preview-03-25", "gemini/gemini-2.5-flash-preview-05-20",
    "gemini/gemini-2.5-flash", "gemini/gemini-2.5-flash-lite-preview-06-17",
    "gemini/gemini-2.5-pro-preview-05-06", "gemini/gemini-2.5-pro-preview-06-05",
    "gemini/gemini-2.5-pro", "gemini/gemini-2.0-flash-exp", "gemini/gemini-2.0-flash",
    "gemini/gemini-2.0-flash-001", "gemini/gemini-2.0-flash-exp-image-generation",
    "gemini/gemini-2.0-flash-lite-001", "gemini/gemini-2.0-flash-lite",
    "gemini/gemini-2.0-flash-preview-image-generation", "gemini/gemini-2.0-flash-lite-preview-02-05",
    "gemini/gemini-2.0-flash-lite-preview", "gemini/gemini-2.0-pro-exp",
    "gemini/gemini-2.0-pro-exp-02-05", "gemini/gemini-exp-1206",
    "gemini/gemini-2.0-flash-thinking-exp-01-21", "gemini/gemini-2.0-flash-thinking-exp",
    "gemini/gemini-2.0-flash-thinking-exp-1219", "gemini/gemini-2.5-flash-preview-tts",
    "gemini/gemini-2.5-pro-preview-tts"
]

# Page Config
st.set_page_config(page_title="Multilingual AI Bible Study Generator", page_icon="🌍", layout="wide")

# Sidebar for API Keys and Model Selection
with st.sidebar:
    st.header("⚙️ API Configuration")
    st.markdown("Enter your API keys below to enable the AI crew.")
    st.session_state.gemini_key = st.text_input("Enter your Google Gemini API Key", type="password")
    st.session_state.serper_key = st.text_input("Enter your Serper API Key", type="password")
    
    st.header("🤖 Model Selection")
    # CORRECTED: Use the new specific list of models
    st.session_state.gemini_model = st.selectbox("Select Gemini Model", GEMINI_MODEL_LIST)

# App Header
st.title("📖 Multilingual AI Bible Study Generator")
st.markdown("Powered by Google Gemini and a team of AI Theologians, Pastors, and Historians.")

# User Input
st.header("1. Select Your Language and Book")
selected_language = st.selectbox("Choose your language:", list(BIBLE_BOOKS_TRANSLATIONS.keys()))
selected_book_translated = st.selectbox("Choose a book to study:", BIBLE_BOOKS_TRANSLATIONS[selected_language])

# Crew Execution
st.header("2. Generate Your Study Guide")
if st.button(f"Create Study Guide for {selected_book_translated}"):
    if not st.session_state.gemini_key or not st.session_state.serper_key:
        st.error("🚨 Please enter your Gemini and Serper API keys in the sidebar to continue.")
    else:
        if "study_guide_content" in st.session_state:
            del st.session_state["study_guide_content"]
        
        book_index = BIBLE_BOOKS_TRANSLATIONS[selected_language].index(selected_book_translated)
        english_book_name = ENGLISH_BOOKS[book_index]

        with st.spinner(f"Your AI Bible Study Team is preparing your guide for '{selected_book_translated}' in {selected_language}..."):
            try:
                bible_study_crew = create_bible_study_crew(
                    bible_book=english_book_name,
                    language=selected_language,
                    selected_model=st.session_state.gemini_model,
                    gemini_api_key=st.session_state.gemini_key,
                    serper_api_key=st.session_state.serper_key
                )
                result = bible_study_crew.kickoff()
                
                output_filename = f'final_study_guide_{selected_language.lower()}.md'
                with open(output_filename, 'r', encoding='utf-8') as file:
                    st.session_state["study_guide_content"] = file.read()
                
                st.success("Your study guide is ready!")
                st.balloons()
            except Exception as e:
                st.error(f"An error occurred: {e}")

# Display and Export Section
if "study_guide_content" in st.session_state:
    st.header("3. Your Custom Study Guide")
    guide_content = st.session_state["study_guide_content"]
    st.markdown(guide_content)
    
    st.header("4. Export Your Guide")
    col1, col2, col3, col4, col5 = st.columns(5)
    filename_base = f"{selected_book_translated.replace(' ', '_')}_study_guide"
    
    with col1:
        st.download_button("⬇️ Download as PDF", markdown_to_pdf(guide_content), f"{filename_base}.pdf", "application/pdf")
    with col2:
        st.download_button("⬇️ Download as DOCX", markdown_to_docx(guide_content), f"{filename_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col3:
        st.download_button("⬇️ Download as MD", guide_content.encode('utf-8'), f"{filename_base}.md", "text/markdown")
    with col4:
        st.download_button("⬇️ Download as HTML", guide_content.encode('utf-8'), f"{filename_base}.html", "text/html")
    with col5:
        st.download_button("⬇️ Download as TXT", guide_content.encode('utf-8'), f"{filename_base}.txt", "text/plain")